from pathlib import Path

from docx import Document

p = Path(__file__).resolve().parent / "DocentDesk_Project_Report.docx"
d = Document(p)
text = "\n".join([para.text for para in d.paragraphs if para.text])
words = len(text.split())
print("File:", p)
print("Paragraphs:", len(d.paragraphs))
print("Approx words:", words)
print("Approx pages (350 wpp):", round(words / 350, 1))
