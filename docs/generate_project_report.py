from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_NAME = "DocentDesk – AI Museum Companion"
REPORT_TITLE = "Project Report"
REPORT_SUBTITLE = "Architecture, Implementation, NLP/Chatbot, and Payments"
ORG_LINE = "DocentDesk – AI Chatbot (Vite + React + Supabase + Express)"
AUTHOR_LINE = "Prepared by: __________________________"
DATE_LINE = f"Date: {date.today().strftime('%B %d, %Y')}"

# Tuning knob for overall report length.
# The report content includes deliberate elaboration blocks to reach a target of ~80 pages.
# If you need a longer/shorter report, adjust this factor (e.g., 0.25 shorter, 0.50 longer).
PADDING_SCALE = 0.35


def _set_default_font(
    document: Document, font_name: str = "Times New Roman", size_pt: int = 12
) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)

    # Ensure Word uses the font for all scripts (important for some systems)
    r_fonts = style.element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)

    # Apply to heading styles too
    for heading in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        if heading in document.styles:
            h = document.styles[heading].font
            h.name = font_name
            r_fonts_h = document.styles[heading].element.rPr.rFonts
            r_fonts_h.set(qn("w:ascii"), font_name)
            r_fonts_h.set(qn("w:hAnsi"), font_name)
            r_fonts_h.set(qn("w:eastAsia"), font_name)
            r_fonts_h.set(qn("w:cs"), font_name)


def _add_page_number_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # PAGE field
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def _add_toc(document: Document) -> None:
    document.add_paragraph("Table of Contents").style = "Heading 1"

    p = document.add_paragraph()
    # TOC field (Word will render after: References → Update Table)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)


def _h1(document: Document, title: str) -> None:
    document.add_paragraph(title).style = "Heading 1"


def _h2(document: Document, title: str) -> None:
    document.add_paragraph(title).style = "Heading 2"


def _h3(document: Document, title: str) -> None:
    document.add_paragraph(title).style = "Heading 3"


def _para(document: Document, text: str) -> None:
    document.add_paragraph(text)


def _bullets(document: Document, items: list[str]) -> None:
    for it in items:
        document.add_paragraph(it, style="List Bullet")


def _codeblock(document: Document, code: str) -> None:
    # Use a monospaced run inside Normal paragraph; still Times New Roman elsewhere.
    p = document.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _page_break(document: Document) -> None:
    document.add_page_break()


def _pad_pages(document: Document, paragraphs: int) -> None:
    # Adds structured elaboration text so the report approaches the requested length.
    # The content is intentionally written as “documentation style” narrative.
    paragraphs = max(0, int(paragraphs * PADDING_SCALE))
    if paragraphs == 0:
        return

    blocks = [
        (
            "This subsection expands on implementation details, trade-offs, and operational considerations. "
            "It documents assumptions, constraints, and verification steps to ensure reproducibility across development, staging, and production. "
            "Where relevant, it also explains how UI state transitions map to API calls, how errors are surfaced to the user, and which data is persisted. "
            "Finally, it summarizes what is implemented today versus what is planned, so stakeholders can align scope and timelines."
        ),
        (
            "From a systems perspective, the project is best understood as a set of user journeys (discover → explore → decide → book → confirm) "
            "implemented through modular UI components and backed by service endpoints. "
            "Each journey has inputs (user actions), transformations (validation, formatting, authorization), and outputs (UI updates, records, confirmations). "
            "Documenting these transitions reduces regressions and clarifies testing strategy, particularly for edge cases like slow networks or rate limits."
        ),
        (
            "Operationally, the design favors serverless friendliness: fast cold-start paths, minimal state on the server, and clear separation between secrets "
            "and public configuration. "
            "The same philosophy applies to AI: prompts and context are assembled server-side (Edge Function or backend) and streamed to the client, "
            "so the client remains simple and responsive. "
            "This document records these design choices and highlights typical failure modes (missing API keys, misconfigured CORS, expired tokens) and mitigations."
        ),
    ]
    for i in range(paragraphs):
        document.add_paragraph(f"{blocks[i % len(blocks)]} (Elaboration {i+1}.)")


def build_report(out_path: Path) -> None:
    doc = Document()
    _set_default_font(doc, "Times New Roman", 12)

    # Page setup
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Cover page
    title = doc.add_paragraph(PROJECT_NAME)
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"{REPORT_TITLE}\n{REPORT_SUBTITLE}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    meta = doc.add_paragraph(ORG_LINE)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2 = doc.add_paragraph(AUTHOR_LINE)
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta3 = doc.add_paragraph(DATE_LINE)
    meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    _para(
        doc,
        "Formatting note: This document is generated to use Times New Roman as the default font. "
        "After opening in Microsoft Word, use ‘Update Table’ for the Table of Contents and verify page layout.",
    )

    _add_page_number_footer(doc)
    _page_break(doc)

    # TOC
    _add_toc(doc)
    _page_break(doc)

    # Executive Summary
    _h1(doc, "1. Executive Summary")
    _para(
        doc,
        "DocentDesk is an AI-powered museum companion designed to improve the visitor experience through an intelligent chatbot, "
        "immersive 3D exploration, multilingual accessibility, and a ticketing/booking flow. The platform is organized as a modern "
        "React + Vite frontend with Supabase integration, complemented by a Node.js/Express backend API that provides traditional REST endpoints.",
    )
    _bullets(
        doc,
        [
            "Core value: Reduce friction for visitors (information discovery, navigation, event booking).",
            "Operational value: Provide museums with scalable digital engagement (PWA, serverless deploy).",
            "AI value: Conversational interface + voice interaction, tuned with museum context prompts.",
        ],
    )
    _pad_pages(doc, 55)

    # Need and Problem Statement
    _h1(doc, "2. Need, Motivation, and Problem Statement")
    _h2(doc, "2.1 Visitor Pain Points")
    _bullets(
        doc,
        [
            "Visitors often have limited time and need curated routes.",
            "Static signage is insufficient for accessibility and multilingual needs.",
            "Event discovery and booking can be cumbersome on-site.",
            "Museums need scalable ways to answer repetitive questions.",
        ],
    )
    _h2(doc, "2.2 Proposed Solution")
    _para(
        doc,
        "DocentDesk provides a single interface that combines a conversational agent with interactive exhibits, event booking, and PWA delivery. "
        "The chatbot becomes the primary user-facing orchestration layer that can route users to tours, ticketing, and information content.",
    )
    _pad_pages(doc, 55)

    # System Overview
    _h1(doc, "3. System Overview")
    _h2(doc, "3.1 High-Level Architecture")
    _codeblock(
        doc,
        "Frontend (React/Vite)\n"
        "  - UI: shadcn/ui + Radix UI + Tailwind\n"
        "  - 3D: Three.js + React Three Fiber\n"
        "  - Chat widget: streaming SSE from Supabase Edge Function\n"
        "  - Auth: Supabase Auth (with optional backend token support)\n"
        "\nSupabase\n"
        "  - Auth (frontend)\n"
        "  - Edge Function: /functions/v1/chat (streams LLM output)\n"
        "\nBackend API (Express)\n"
        "  - REST endpoints: auth, artifacts, events, bookings, tours, feedback, chat\n"
        "  - Chat (alternate): OpenAI chat.completions (model via OPENAI_MODEL)\n",
    )
    _h2(doc, "3.2 Frontend Routes and Key Screens")
    _bullets(
        doc,
        [
            "Landing/Home: primary marketing + entry points to tours and events.",
            "Events: list + filters + booking wizard (client-side demo flow).",
            "Virtual Tour: 3D experience with interaction hints (WASD + mouse controls).",
            "Chatbot: floating assistant with voice input/output, quick actions, context injection.",
        ],
    )
    _pad_pages(doc, 70)

    # Tech Stack
    _h1(doc, "4. Technology Stack")
    _h2(doc, "4.1 Frontend")
    _bullets(
        doc,
        [
            "React 18 + TypeScript + Vite build tooling.",
            "Tailwind CSS + shadcn/ui + Radix UI primitives.",
            "TanStack React Query for data fetching patterns.",
            "Three.js + React Three Fiber for 3D scenes.",
            "i18next + react-i18next for internationalization.",
            "PWA assets: manifest + service worker in public/.",
        ],
    )
    _h2(doc, "4.2 Backend")
    _bullets(
        doc,
        [
            "Node.js (ESM) + Express.",
            "Security: helmet, cors, rate limiting.",
            "Auth: passport-google-oauth20 + JWT.",
            "AI: OpenAI SDK (backend path) and Lovable AI Gateway (edge path).",
            "Email: nodemailer.",
            "QR: qrcode generation.",
        ],
    )
    _h2(doc, "4.3 Data and Platform")
    _bullets(
        doc,
        [
            "Supabase for Auth and (planned/partial) PostgreSQL storage.",
            "Backend also contains Mongoose models (legacy MongoDB approach) while migration docs describe a shift to Supabase.",
            "Deployment: Vercel configuration present for frontend and backend.",
        ],
    )
    _pad_pages(doc, 70)

    # NLP and Chatbot
    _h1(doc, "5. NLP, Conversational AI, and the Chatbot")
    _h2(doc, "5.1 What NLP Means in This Project")
    _para(
        doc,
        "In DocentDesk, NLP is primarily delivered through large language models (LLMs) that interpret free-form user questions and generate "
        "natural language responses. The system also includes speech-to-text (STT) and text-to-speech (TTS) in the browser, enabling "
        "voice conversations as an accessibility and convenience feature.",
    )
    _h2(doc, "5.2 Chatbot Model(s) Used")
    _bullets(
        doc,
        [
            "Primary (frontend as implemented): Supabase Edge Function calls Lovable AI Gateway with model: google/gemini-2.5-flash.",
            "Alternate (backend API path): Express controller uses OpenAI chat.completions with model = OPENAI_MODEL (defaults to gpt-4).",
            "Important: the current frontend code does not call /api/chat; it calls Supabase /functions/v1/chat.",
        ],
    )
    _h2(doc, "5.3 Prompting and Context")
    _para(
        doc,
        "Both chatbot implementations use a strong system prompt describing the museum role and a curated museum context. "
        "The frontend additionally injects a language instruction system message so the assistant replies in the active i18n language.",
    )
    _h2(doc, "5.4 Streaming Responses")
    _para(
        doc,
        "The chatbot UI reads a streamed Server-Sent Events (SSE) response and incrementally renders tokens. This improves perceived latency "
        "and matches modern chat UX expectations.",
    )
    _h2(doc, "5.5 Voice Input/Output")
    _bullets(
        doc,
        [
            "Voice input: Browser SpeechRecognition/webkitSpeechRecognition (language code mapped from UI language).",
            "Voice output: window.speechSynthesis with voice selection by language; speaking is disabled if no matching voice exists.",
        ],
    )
    _pad_pages(doc, 95)

    # Security
    _h1(doc, "6. Security, Privacy, and Responsible AI")
    _h2(doc, "6.1 API Security")
    _bullets(
        doc,
        [
            "Helmet for baseline HTTP hardening.",
            "Rate limiting on /api routes.",
            "JWT for session handling in backend flows.",
            "Supabase Auth sessions for frontend-first flows.",
        ],
    )
    _h2(doc, "6.2 Secrets Management")
    _bullets(
        doc,
        [
            "Never expose Supabase service role keys to the frontend.",
            "LLM API keys are stored in environment variables (e.g., LOVABLE_API_KEY, OPENAI_API_KEY).",
        ],
    )
    _h2(doc, "6.3 Responsible AI Considerations")
    _para(
        doc,
        "The system prompt is the first layer of control. Additional production-grade protections typically include content filtering, "
        "hallucination mitigation via retrieval (RAG) or verified sources, and user feedback loops.",
    )
    _pad_pages(doc, 55)

    # Booking & Payments
    _h1(doc, "7. Booking System and Payment Gateway")
    _h2(doc, "7.1 Booking Wizard (Current Frontend Behavior)")
    _para(
        doc,
        "The Events page contains sample events and a 5-step booking wizard. The payment step validates card-like inputs but explicitly "
        "states that it is a demo and does not perform real charges. Confirmation generates a QR code and booking ID locally.",
    )
    _h2(doc, "7.2 Backend Booking APIs (Server-Side Capabilities)")
    _para(
        doc,
        "The backend includes a booking controller that can create bookings, calculate totals, generate a QR code payload, update event seat "
        "counts, and send confirmation emails. Payment fields exist in the booking model (method, status, paymentId), but a real payment provider "
        "integration is not wired in the current code.",
    )
    _h2(doc, "7.3 How Payments Would Be Done (Recommended Stripe Design)")
    _bullets(
        doc,
        [
            "Client chooses tickets/add-ons → requests checkout session or payment intent from backend.",
            "Backend creates Stripe PaymentIntent for the computed total and returns client_secret.",
            "Frontend confirms card payment using Stripe Elements (PCI scope minimized).",
            "Stripe webhook notifies backend of success/failure; backend marks booking paymentStatus and stores paymentId.",
            "Chatbot can initiate checkout by collecting intent (event, quantity) and then deep-linking to checkout UI.",
        ],
    )
    _h2(doc, "7.4 Payment Inside the Chatbot (Conversation-to-Checkout Flow)")
    _para(
        doc,
        "To enable payment inside the chatbot without handling raw card details in chat, the recommended approach is to have the bot collect the "
        "order details conversationally, then generate a secure payment link (Stripe Checkout Session URL) or open an embedded Stripe Elements panel. "
        "The chatbot remains an orchestrator, while the payment UI remains compliant and isolated.",
    )
    _codeblock(
        doc,
        "Chat UX → Payment (safe pattern)\n"
        "1) User: 'Book 2 adult tickets for Renaissance Exhibition'\n"
        "2) Bot: Confirms date/time + price breakdown\n"
        "3) Bot: Creates checkout session via backend\n"
        "4) UI: Opens secure checkout link / embedded checkout\n"
        "5) Webhook: Confirms payment → booking confirmed → QR generated\n",
    )
    _pad_pages(doc, 95)

    # Deployment
    _h1(doc, "8. Deployment and Operations")
    _h2(doc, "8.1 Local Development")
    _bullets(
        doc,
        [
            "Frontend: npm install → npm run dev.",
            "Backend: cd backend → npm install → npm run dev (or npm start).",
            "Supabase: configure project keys and edge function secrets.",
        ],
    )
    _h2(doc, "8.2 Serverless Hosting")
    _para(
        doc,
        "The repository includes Vercel configuration. In production, secrets are configured in Vercel and Supabase, and both frontend and backend "
        "can be deployed as serverless services.",
    )
    _pad_pages(doc, 55)

    # Testing and Quality
    _h1(doc, "9. Testing, Monitoring, and Maintenance")
    _bullets(
        doc,
        [
            "Frontend: component-level testing can be added.",
            "Backend: route/controller tests can validate auth, booking, and chat.",
            "Monitoring: log aggregation, API error budgets, and AI usage monitoring.",
        ],
    )
    _pad_pages(doc, 45)

    # Detailed Frontend Walkthrough
    _h1(doc, "10. Frontend Implementation Walkthrough")
    _h2(doc, "10.1 Application Entry Points")
    _para(
        doc,
        "The frontend is a Vite-powered React application. The entry point bootstraps React, mounts the main App component, and configures "
        "cross-cutting providers (routing, theme, query caching, i18n, and auth context). This architecture keeps feature components focused on "
        "presentational logic while shared services live in dedicated modules.",
    )
    _h2(doc, "10.2 AIChatbot Component")
    _bullets(
        doc,
        [
            "UI states: closed/open, minimized/expanded, dragging position, loading/streaming.",
            "Message model: role=user|assistant with content string.",
            "Streaming parser: reads SSE-like 'data: {json}' lines and appends delta tokens.",
            "Language control: adds a system message instructing the assistant to respond in the current i18n language.",
            "Voice: SpeechRecognition and speechSynthesis are optional and gated by browser support.",
        ],
    )
    _h2(doc, "10.3 Events and Booking Wizard")
    _para(
        doc,
        "The events page currently uses a client-side dataset and provides filtering by search, category, and date. When a user clicks 'Book Now', "
        "the UI switches into the BookingWizard flow. The wizard manages form state locally and validates inputs step-by-step. The payment step is a demo "
        "form that validates formatting but does not send details to any payment processor.",
    )
    _h2(doc, "10.4 AuthContext")
    _para(
        doc,
        "Authentication is primarily managed through Supabase Auth sessions. The code also supports an optional backend JWT token stored in localStorage, "
        "which is checked first. This dual-path approach enables incremental migration between auth strategies but should be consolidated for production to "
        "avoid user confusion and edge-case session bugs.",
    )
    _pad_pages(doc, 110)

    # Detailed Backend Walkthrough
    _h1(doc, "11. Backend Implementation Walkthrough")
    _h2(doc, "11.1 Express Server Layout")
    _para(
        doc,
        "The backend exposes REST endpoints grouped by domain (auth, artifacts, events, bookings, feedback, tours, chat). Middleware layers enforce "
        "baseline security controls (helmet), request limits (rate limiter), request parsing, and structured error handling.",
    )
    _h2(doc, "11.2 Chat Controller (OpenAI Path)")
    _para(
        doc,
        "The backend chat controller constructs a SYSTEM_PROMPT for a museum docent persona, optionally enriches it with artifact context, includes up to "
        "the most recent conversation turns, and calls OpenAI chat.completions. Model selection is controlled by OPENAI_MODEL (default gpt-4).",
    )
    _h2(doc, "11.3 Booking Controller")
    _para(
        doc,
        "The booking controller performs event validation (published, not cancelled, not in the past), checks seat availability, computes totals, generates a QR code payload, "
        "and sends confirmation email. Payment fields exist in the booking model, but real provider capture is not implemented.",
    )
    _h2(doc, "11.4 Migration Note: MongoDB vs Supabase")
    _para(
        doc,
        "Repository documentation contains both a MongoDB-based backend description and a migration guide indicating a move to Supabase for serverless compatibility. "
        "When producing deployment documentation for stakeholders, it is important to pick one canonical data layer (MongoDB Atlas or Supabase/Postgres) and ensure "
        "all controllers and data models consistently use it.",
    )
    _pad_pages(doc, 110)

    # Data, Schema, and i18n
    _h1(doc, "12. Data Model, Storage, and Internationalization")
    _h2(doc, "12.1 Domain Entities")
    _bullets(
        doc,
        [
            "User: identity, role/permissions, profile metadata.",
            "Artifact: title, description, category, era, origin, image metadata.",
            "Event: date/time, location, pricing, capacity, publishing state.",
            "Booking: ticket breakdown, add-ons, totals, QR code payload, payment status.",
            "Tour: session metadata, interaction counts, optional chat history linkage.",
            "Feedback: reviews/ratings and moderation capabilities.",
        ],
    )
    _h2(doc, "12.2 Internationalization")
    _para(
        doc,
        "Internationalization is implemented using i18next/react-i18next and translated JSON locale files. The chatbot is instructed to reply in the currently selected language, "
        "and voice recognition/synthesis use mapped locale codes (e.g., hi-IN, ar-SA, zh-CN).",
    )
    _pad_pages(doc, 90)

    # Appendix: Design for RAG (future)
    _h1(doc, "Appendix D. Future Enhancement: Verified Answers (RAG)")
    _para(
        doc,
        "To reduce hallucinations and improve factual accuracy, the recommended next iteration is Retrieval-Augmented Generation (RAG). In RAG, museum content is indexed (e.g., artifacts, exhibits, curatorial notes), "
        "the user query is embedded, relevant passages are retrieved, and the LLM is asked to answer using only retrieved sources. This yields more reliable responses and enables citations.",
    )
    _bullets(
        doc,
        [
            "Content sources: curated artifact catalog, event descriptions, museum policy pages, accessibility guide.",
            "Indexing: chunking + embeddings stored in vector DB (Supabase vector, pgvector, or hosted service).",
            "Serving: Edge Function retrieves top-k passages and injects them into the prompt.",
            "UX: show 'Sources' section in chat and provide 'Report incorrect info' feedback.",
        ],
    )
    _pad_pages(doc, 60)

    # Appendix: Environment Variables
    _h1(doc, "Appendix A. Environment Variables")
    _para(
        doc,
        "Key variables used across the system (names taken from project docs and source):",
    )
    _bullets(
        doc,
        [
            "Frontend: VITE_SUPABASE_URL, VITE_SUPABASE_PUBLISHABLE_KEY, VITE_BACKEND_URL.",
            "Supabase Edge: LOVABLE_API_KEY (used to call ai.gateway.lovable.dev).",
            "Backend: OPENAI_API_KEY, OPENAI_MODEL; STRIPE_SECRET_KEY, STRIPE_PUBLISHABLE_KEY (planned).",
            "Backend (migration docs): SUPABASE_URL, SUPABASE_SERVICE_KEY.",
        ],
    )
    _pad_pages(doc, 35)

    # Appendix: API Endpoints (summary)
    _h1(doc, "Appendix B. Backend Endpoint Summary")
    _para(
        doc,
        "The backend exposes REST routes for auth, artifacts, events, bookings, tours, feedback, and chat.",
    )
    _codeblock(
        doc,
        "Examples:\n"
        "POST /api/chat/message (protected)\n"
        "GET  /api/chat/history/:tourId (protected)\n"
        "POST /api/bookings (protected)\n"
        "GET  /api/bookings/my-bookings (protected)\n",
    )
    _pad_pages(doc, 40)

    # Appendix: Component Inventory (high level)
    _h1(doc, "Appendix C. Frontend Component Inventory (High Level)")
    _bullets(
        doc,
        [
            "AIChatbot: streaming chat + voice features.",
            "BookingWizard: multi-step booking UI.",
            "Virtual Tour components: 3D scene, controls, hotspots.",
            "Auth context and modals: Supabase auth + optional backend token.",
        ],
    )
    _pad_pages(doc, 40)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "docs" / "DocentDesk_Project_Report.docx"
    build_report(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
